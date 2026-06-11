import base64
import io
import os
import re
import requests
import shutil
import subprocess
import tempfile
import uuid
from pathlib import Path

import fitz
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from app.config import settings


DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
CONVERTIBLE_TO_DOCX_MIME_TYPES = {
    "application/msword",
    DOCX_MIME,
    "application/pdf",
    "application/vnd.oasis.opendocument.text",
    "application/rtf",
    "text/rtf",
}


def storage_path(filename: str) -> str:
    return os.path.join(settings.STORAGE_ROOT, filename)


def _soffice() -> str | None:
    return shutil.which("libreoffice") or shutil.which("soffice")


def _run_libreoffice_convert(input_path: str, output_ext: str, timeout: int = 120) -> str | None:
    soffice = _soffice()
    if not soffice:
        return None

    with tempfile.TemporaryDirectory() as tmpdir:
        profile = os.path.join(tmpdir, "lo-profile")
        os.makedirs(profile, exist_ok=True)
        subprocess.run(
            [
                soffice,
                f"-env:UserInstallation=file://{profile}",
                "--headless",
                "--convert-to",
                output_ext,
                "--outdir",
                tmpdir,
                input_path,
            ],
            check=True,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            timeout=timeout,
        )
        matches = [f for f in os.listdir(tmpdir) if f.lower().endswith(f".{output_ext.lower()}")]
        if not matches:
            return None
        output_path = os.path.join(settings.STORAGE_ROOT, f"{uuid.uuid4()}.{output_ext.lower()}")
        shutil.copy(os.path.join(tmpdir, matches[0]), output_path)
        return os.path.basename(output_path)


def _pdf_has_selectable_text(input_path: str) -> bool:
    with fitz.open(input_path) as doc:
        sample_pages = min(len(doc), 5)
        text = "\n".join((doc[idx].get_text("text") or "").strip() for idx in range(sample_pages))
    return len(text.strip()) >= 80


def _clean_pdf_line(text: str) -> str:
    value = " ".join((text or "").replace(" ", " ").split())
    value = value.replace("…..", ".....")
    value = value.replace("…", "...")
    value = __import__("re").sub(r"\.{3,}\s*\d{0,4}$", "", value).strip()
    return value


def _line_is_page_number(text: str, y0: float, page_height: float) -> bool:
    value = (text or "").strip()
    return value.isdigit() and y0 > page_height - 80


def _line_is_toc_title(text: str) -> bool:
    return bool(__import__("re").search(r"(?i)\b(table of contents|contents|yaliyomo|okuqukethwe|zviri\s+mukati)\b", text or ""))


def _line_is_chapter_heading(text: str) -> bool:
    return bool(__import__("re").match(
        r"(?i)^(chapter|sura(?:\s+ya)?|isahluko|chitsauko|hoofstuk|section|sehemu|isigaba)\b",
        (text or "").strip(),
    ))


def _line_is_tail_promo_start(text: str) -> bool:
    return bool(__import__("re").match(
        r"(?i)^(if\s+you\s+are\s+looking\s+for|iwe\s+unatafuta|uma\s+ufuna|kungakhathaliseki\s+ukuthi\s+ufuna)\b",
        (text or "").strip(),
    ))


def _line_is_legal_credit(text: str) -> bool:
    return bool(__import__("re").search(
        r"(?i)(scripture quotations.*copyright|copyright|©|thomas nelson|international bible society|"
        r"tyndale house|used by permission|all rights reserved|no part of this publication|publisher)",
        text or "",
    ))


def _line_is_allcaps_heading(text: str) -> bool:
    value = (text or "").strip()
    if not value or len(value) > 140:
        return False
    letters = [ch for ch in value if ch.isalpha()]
    if not letters:
        return False
    upper_ratio = sum(1 for ch in letters if ch.upper() == ch) / len(letters)
    return upper_ratio >= 0.78 and not value.endswith((".", ",", ";"))


def _extract_pdf_image_blocks(page) -> list[dict]:
    items = []
    for block in page.get_text("dict", sort=True).get("blocks", []):
        if block.get("type") != 1:
            continue
        image_bytes = block.get("image")
        if not image_bytes:
            continue
        bbox = block.get("bbox") or (0, 0, page.rect.width, page.rect.height)
        items.append({
            "kind": "image",
            "bbox": bbox,
            "image": image_bytes,
            "ext": block.get("ext") or "png",
        })
    return items


def _add_pdf_image_paragraph(document: Document, image_item: dict, page_width: float) -> None:
    image_bytes = image_item.get("image")
    if not image_bytes:
        return
    bbox = image_item.get("bbox") or (0, 0, page_width, page_width)
    width_ratio = max(0.1, min(1.0, (bbox[2] - bbox[0]) / max(page_width, 1)))
    width_inches = max(1.0, min(6.6, 6.6 * width_ratio))
    para = document.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_after = Pt(4)
    run = para.add_run()
    try:
        run.add_picture(io.BytesIO(image_bytes), width=Inches(width_inches))
    except Exception:
        return


def _extract_pdf_lines(page) -> list[dict]:
    rows = []
    for block_index, block in enumerate(page.get_text("dict", sort=True).get("blocks", [])):
        if block.get("type") != 0:
            continue
        for line in block.get("lines", []):
            spans = [span for span in line.get("spans", []) if (span.get("text") or "").strip()]
            if not spans:
                continue
            raw = "".join(span.get("text", "") for span in spans)
            text = _clean_pdf_line(raw)
            if not text:
                continue
            bbox = line.get("bbox") or spans[0].get("bbox") or (0, 0, 0, 0)
            if _line_is_page_number(text, bbox[1], page.rect.height):
                continue
            bold_chars = sum(len(span.get("text", "")) for span in spans if "bold" in (span.get("font") or "").lower())
            total_chars = sum(len(span.get("text", "")) for span in spans) or 1
            max_size = max(float(span.get("size") or 0) for span in spans)
            rows.append({
                "text": text,
                "bbox": bbox,
                "bold": bold_chars / total_chars >= 0.45,
                "size": max_size,
                "centered": abs(((bbox[0] + bbox[2]) / 2) - (page.rect.width / 2)) < page.rect.width * 0.12,
                "block_index": block_index,
            })
    rows.sort(key=lambda item: (round(item["bbox"][1], 1), item["bbox"][0]))
    return rows


def _configure_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)
    normal = document.styles["Normal"]
    normal.font.name = "Liberation Serif"
    normal.font.size = Pt(11)


def _add_pdf_line_paragraph(document: Document, line: dict, page_index: int, allow_promo_page_break: bool = True) -> None:
    text = line["text"]
    if allow_promo_page_break and _line_is_tail_promo_start(text):
        document.add_page_break()
    para = document.add_paragraph()
    para.paragraph_format.space_after = Pt(3)
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.line_spacing = 1.05
    is_chapter = _line_is_chapter_heading(text)
    is_toc_title = _line_is_toc_title(text)
    is_title_page = page_index == 1
    is_heading = is_chapter or is_toc_title or _line_is_allcaps_heading(text)
    if is_chapter or is_toc_title or is_title_page:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = para.add_run(text)
    run.font.name = "Liberation Serif"
    run.font.size = Pt(12 if is_heading else 11)
    run.bold = bool(line.get("bold") or is_heading or is_title_page or _line_is_legal_credit(text))


def _line_is_list_or_table_item(text: str) -> bool:
    value = (text or "").strip()
    return bool(re.match(r"^(?:[•\-*]|\d+[.)]|[A-Za-z][.)])\s+", value))


def _line_ends_sentence(text: str) -> bool:
    return bool(re.search(r"[.!?;:)\']$", (text or "").strip()))


def _pdf_line_starts_new_paragraph(previous: dict | None, current: dict, page_width: float) -> bool:
    if previous is None:
        return True
    prev_text = previous.get("text", "")
    text = current.get("text", "")
    if not prev_text or not text:
        return True
    if _line_is_chapter_heading(text) or _line_is_toc_title(text) or _line_is_allcaps_heading(text):
        return True
    if _line_is_legal_credit(text) or _line_is_tail_promo_start(text):
        return True
    if previous.get("block_index") != current.get("block_index"):
        return True
    if _line_is_list_or_table_item(text):
        return True
    if previous.get("bold") != current.get("bold") and (previous.get("bold") or current.get("bold")):
        return True
    prev_bbox = previous.get("bbox") or (0, 0, 0, 0)
    bbox = current.get("bbox") or (0, 0, 0, 0)
    prev_height = max(1.0, float(prev_bbox[3] - prev_bbox[1]))
    vertical_gap = float(bbox[1] - prev_bbox[3])
    if vertical_gap > prev_height * 0.70:
        return True
    left_delta = abs(float(bbox[0] - prev_bbox[0]))
    if left_delta > max(18.0, page_width * 0.035) and _line_ends_sentence(prev_text):
        return True
    return False


def _flush_pdf_text_group(document: Document, group: list[dict], page_index: int, allow_promo_page_break: bool = True) -> None:
    if not group:
        return
    first = group[0]
    merged = " ".join((item.get("text") or "").strip() for item in group if (item.get("text") or "").strip())
    if not merged:
        return
    styled_line = dict(first)
    styled_line["text"] = merged
    _add_pdf_line_paragraph(document, styled_line, page_index, allow_promo_page_break=allow_promo_page_break)


def _convertapi_base_url() -> str:
    region = (settings.CONVERTAPI_REGION or "v2").strip().lower()
    allowed = {"v2", "eu-v2", "uk-v2", "us-v2", "ca-v2", "as-v2", "au-v2"}
    if region not in allowed:
        region = "v2"
    return f"https://{region}.convertapi.com"


def _convertapi_pdf_to_docx(input_path: str) -> str | None:
    token = (settings.CONVERTAPI_TOKEN or "").strip()
    if not token:
        return None

    output_path = os.path.join(settings.STORAGE_ROOT, f"{uuid.uuid4()}.docx")
    url = f"{_convertapi_base_url()}/convert/pdf/to/docx"
    headers = {
        "Authorization": f"Bearer {token}",
        "Accept": "application/octet-stream, application/json",
    }
    data = {
        "StoreFile": "false",
        "OcrMode": "never",
        "Wysiwyg": "false",
        "Timeout": "600",
    }
    with open(input_path, "rb") as source_file:
        response = requests.post(
            url,
            headers=headers,
            files={"File": (os.path.basename(input_path), source_file, "application/pdf")},
            data=data,
            timeout=660,
        )
    response.raise_for_status()

    content_type = (response.headers.get("content-type") or "").lower()
    if "application/json" in content_type:
        payload = response.json()
        files = payload.get("Files") or []
        if not files:
            raise ValueError("ConvertAPI response did not include converted files")
        first_file = files[0]
        if first_file.get("FileData"):
            content = base64.b64decode(first_file["FileData"])
        elif first_file.get("Url"):
            file_response = requests.get(first_file["Url"], timeout=660)
            file_response.raise_for_status()
            content = file_response.content
        else:
            raise ValueError("ConvertAPI file response did not include FileData or Url")
    else:
        content = response.content

    if not content or len(content) < 500:
        raise ValueError("ConvertAPI returned an empty or invalid DOCX")
    if not content.startswith(b"PK"):
        raise ValueError("ConvertAPI response is not a DOCX zip file")

    with open(output_path, "wb") as out_file:
        out_file.write(content)
    return os.path.basename(output_path)


def _structured_selectable_pdf_to_docx(input_path: str) -> str:
    """Create a DOCX from selectable PDF text with workbook-aware page/style structure."""
    if not _pdf_has_selectable_text(input_path):
        raise ValueError("PDF does not contain enough selectable text to convert without OCR")

    document = Document()
    _configure_document(document)
    with fitz.open(input_path) as pdf:
        for page_index, page in enumerate(pdf):
            lines = _extract_pdf_lines(page)
            images = _extract_pdf_image_blocks(page)
            page_area = page.rect.width * page.rect.height
            large_images = [img for img in images if ((img["bbox"][2] - img["bbox"][0]) * (img["bbox"][3] - img["bbox"][1])) >= page_area * 0.35]

            # Preserve image-based covers as images. Rebuilding them as text loses logos/artwork.
            if page_index == 0 and large_images:
                for image in large_images:
                    _add_pdf_image_paragraph(document, image, page.rect.width)
            else:
                line_items = [{"kind": "text", **line} for line in lines]
                image_items = images
                items = sorted(line_items + image_items, key=lambda item: (round(item["bbox"][1], 1), item["bbox"][0]))
                text_index = 0
                text_group = []
                previous_text_item = None
                for item in items:
                    if item["kind"] == "image":
                        _flush_pdf_text_group(document, text_group, page_index, allow_promo_page_break=text_index > len(text_group))
                        text_group = []
                        previous_text_item = None
                        _add_pdf_image_paragraph(document, item, page.rect.width)
                    else:
                        if _pdf_line_starts_new_paragraph(previous_text_item, item, page.rect.width):
                            _flush_pdf_text_group(document, text_group, page_index, allow_promo_page_break=text_index > len(text_group))
                            text_group = [item]
                        else:
                            text_group.append(item)
                        previous_text_item = item
                        text_index += 1
                _flush_pdf_text_group(document, text_group, page_index, allow_promo_page_break=text_index > len(text_group))

                if page_index == 0 and not lines and not images:
                    para = document.add_paragraph()
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = para.add_run("TEAM IMPACT CHRISTIAN UNIVERSITY")
                    run.bold = True
                    run.font.name = "Liberation Serif"
                    run.font.size = Pt(18)
            if page_index < len(pdf) - 1:
                document.add_page_break()

    output_path = os.path.join(settings.STORAGE_ROOT, f"{uuid.uuid4()}.docx")
    document.save(output_path)
    return os.path.basename(output_path)


def _docx_has_enough_extractable_text(file_name: str) -> bool:
    try:
        from app.services.docx_translation_service import extract_docx_translation_text
        with open(storage_path(file_name), "rb") as docx_file:
            text = extract_docx_translation_text(docx_file.read())
        return len(text.strip()) >= 500
    except Exception:
        return False


def _selectable_pdf_to_docx(input_path: str) -> str:
    """Fallback PDF-to-DOCX for text PDFs. Preserves text order, not exact PDF layout."""
    if not _pdf_has_selectable_text(input_path):
        raise ValueError("PDF does not contain enough selectable text to convert without OCR")

    document = Document()
    with fitz.open(input_path) as pdf:
        for page_index, page in enumerate(pdf):
            if page_index:
                document.add_page_break()
            blocks = page.get_text("blocks", sort=True)
            for block in blocks:
                text = (block[4] or "").strip()
                if not text:
                    continue
                for paragraph in text.splitlines():
                    paragraph = paragraph.strip()
                    if paragraph:
                        document.add_paragraph(paragraph)

    output_path = os.path.join(settings.STORAGE_ROOT, f"{uuid.uuid4()}.docx")
    document.save(output_path)
    return os.path.basename(output_path)


def normalize_upload_to_docx(file_path: str, mime_type: str) -> tuple[str | None, str | None]:
    """
    Normalize uploaded book documents to DOCX for the structure-preserving renderer.

    OCR hook intentionally disabled for now:
    # if mime_type == "application/pdf" and not _pdf_has_selectable_text(input_path):
    #     return ocr_pdf_to_docx(input_path), None
    """
    input_path = storage_path(file_path)
    if not os.path.exists(input_path):
        return None, f"Source file not found: {file_path}"

    if mime_type == DOCX_MIME or file_path.lower().endswith(".docx"):
        return file_path, None

    try:
        if mime_type == "application/pdf" or file_path.lower().endswith(".pdf"):
            convertapi_error = None
            try:
                converted = _convertapi_pdf_to_docx(input_path)
                if converted and _docx_has_enough_extractable_text(converted):
                    return converted, None
            except Exception as exc:
                convertapi_error = str(exc)
            converted = _structured_selectable_pdf_to_docx(input_path)
            return converted, convertapi_error

        converted = _run_libreoffice_convert(input_path, "docx")
        if converted:
            return converted, None
        return None, f"LibreOffice could not convert {Path(file_path).suffix or mime_type} to DOCX"
    except Exception as exc:
        return None, str(exc)


def compress_pdf_bytes(pdf_bytes: bytes, quality: str = "ebook") -> bytes:
    """Compress PDF with Ghostscript. Returns original bytes if compression fails or grows."""
    gs = shutil.which("gs")
    if not gs or not pdf_bytes:
        return pdf_bytes

    quality = quality if quality in {"screen", "ebook", "printer", "prepress"} else "ebook"
    with tempfile.TemporaryDirectory() as tmpdir:
        input_path = os.path.join(tmpdir, "input.pdf")
        output_path = os.path.join(tmpdir, "output.pdf")
        with open(input_path, "wb") as f:
            f.write(pdf_bytes)
        try:
            subprocess.run(
                [
                    gs,
                    "-sDEVICE=pdfwrite",
                    "-dCompatibilityLevel=1.4",
                    f"-dPDFSETTINGS=/{quality}",
                    "-dNOPAUSE",
                    "-dQUIET",
                    "-dBATCH",
                    f"-sOutputFile={output_path}",
                    input_path,
                ],
                check=True,
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
                timeout=120,
            )
            with open(output_path, "rb") as f:
                compressed = f.read()
            return compressed if compressed and len(compressed) < len(pdf_bytes) else pdf_bytes
        except Exception:
            return pdf_bytes
