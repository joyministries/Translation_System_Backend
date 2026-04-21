import os
import logging

from docx import Document
from openpyxl import load_workbook
from openpyxl.styles import Font, Alignment
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.units import inch
import io

from app.config import settings
from app.utils.file_utils import get_file_path

logger = logging.getLogger(__name__)


def extract_text_from_docx(file_path: str) -> str:
    full_path = get_file_path(file_path)
    doc = Document(full_path)
    return "\n".join([para.text for para in doc.paragraphs])


def extract_text_from_doc(file_path: str) -> str:
    import subprocess

    full_path = get_file_path(file_path)
    result = subprocess.run(
        ["catdoc", full_path], capture_output=True, text=True, timeout=60
    )

    if result.returncode == 0:
        return result.stdout
    logger.warning(f"catdoc failed: {result.stderr}")
    return ""


def translate_excel_from_json(original_file_path: str, translated_json: str) -> bytes:
    import json

    full_path = get_file_path(original_file_path)
    wb = load_workbook(full_path)

    try:
        translations = json.loads(translated_json)
        sheets_data = translations.get("sheets") or translations.get("makaratasi") or {}
        original_names = translations.get("sheet_names") or []
        translated_names = translations.get("translated_sheet_names") or []

        # sheets_data keys are translated names — map back to original workbook sheet names
        # original_names[i] -> workbook sheet, translated_names[i] -> sheets_data key
        if (
            original_names
            and translated_names
            and len(original_names) == len(translated_names)
        ):
            name_map = dict(zip(original_names, translated_names))
        else:
            # fallback: keys match workbook names directly
            name_map = {n: n for n in wb.sheetnames}

        rows_updated = 0
        for orig_name in wb.sheetnames:
            translated_key = name_map.get(orig_name, orig_name)
            rows = sheets_data.get(translated_key) or sheets_data.get(orig_name) or []
            if not rows:
                continue
            ws = wb[orig_name]
            for row_idx, row_data in enumerate(rows, start=1):
                for col_idx, value in enumerate(row_data, start=1):
                    if row_idx <= ws.max_row and col_idx <= ws.max_column:
                        if value and str(value).strip():
                            try:
                                ws.cell(row=row_idx, column=col_idx).value = value
                                rows_updated += 1
                            except AttributeError:
                                pass  # skip merged cells

        # Rename sheets to translated names
        for orig_name, trans_name in name_map.items():
            if orig_name in wb.sheetnames and trans_name and trans_name != orig_name:
                wb[orig_name].title = trans_name

        logger.info(f"Updated {rows_updated} cells in Excel")

    except Exception as e:
        logger.warning(f"Could not apply translations to Excel: {e}")

    output = io.BytesIO()
    wb.save(output)
    output.seek(0)
    return output.getvalue()


def generate_translated_excel(translated_text: str, file_path: str) -> bytes | None:
    import json
    import re

    full_path = get_file_path(file_path)
    try:
        wb = load_workbook(full_path)
    except:
        return None

    try:
        # Fix malformed JSON (missing values before ])
        fixed_text = re.sub(r",\]", ', ""]', translated_text)
        data = json.loads(fixed_text)

        if isinstance(data, dict) and ("sheets" in data or "makaratasi" in data):
            sheets_data = data.get("sheets") or data.get("makaratasi") or {}

            # Map English sheet names to Swahili keys
            sheet_names_map = (
                data.get("majina_ya_karatasi") or data.get("sheet_names") or []
            )
            if sheet_names_map:
                swahili_keys = list(sheets_data.keys())
                if len(swahili_keys) == len(wb.sheetnames):
                    mapped_data = {}
                    for eng_name, swahili_key in zip(wb.sheetnames, swahili_keys):
                        if swahili_key in sheets_data:
                            mapped_data[eng_name] = sheets_data[swahili_key]
                    sheets_data = mapped_data

            for sheet_name in wb.sheetnames:
                if sheet_name in sheets_data:
                    ws = wb[sheet_name]
                    rows = sheets_data[sheet_name]
                    for r_idx, row in enumerate(rows, start=1):
                        for c_idx, val in enumerate(row, start=1):
                            if r_idx <= ws.max_row and c_idx <= ws.max_column:
                                ws.cell(row=r_idx, column=c_idx).value = val

            output = io.BytesIO()
            wb.save(output)
            output.seek(0)
            return output.getvalue()
    except Exception as e:
        logger.warning(f"Could not parse translation: {e}")

    return None


def create_translated_pdf(text: str) -> bytes:
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter)
    styles = getSampleStyleSheet()
    story = []

    paragraphs = text.split("\n")
    for para in paragraphs:
        if para.strip():
            story.append(Paragraph(para, styles["Normal"]))
            story.append(Spacer(1, 0.2 * inch))

    doc.build(story)
    buffer.seek(0)
    return buffer.getvalue()


def create_translated_docx(text: str, cover_text: str = None) -> bytes:
    from docx.shared import Pt

    doc = Document()

    # Add cover page first if provided
    if cover_text:
        for para in cover_text.split("\n"):
            if para.strip():
                p = doc.add_paragraph(para)
                # Make cover text bold
                for run in p.runs:
                    run.font.bold = True
                    run.font.size = Pt(14)
        # Add page break after cover
        doc.add_page_break()

    # Add translated content
    for para in text.split("\n"):
        if para.strip():
            doc.add_paragraph(para)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()
