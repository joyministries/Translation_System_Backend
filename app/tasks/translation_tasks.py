import json
import logging
import os
import uuid
from datetime import datetime

import requests
from celery import Task

from app.tasks.celery_app import celery_app
from app.database import SessionLocal
from app.models import Language
from app.utils.text_chunker import chunk_text, merge_chunks
from app.config import settings

logger = logging.getLogger(__name__)

TRANSLATION_ENGINE = "google_translate"


def translate_chunk(text: str, source_lang: str, target_lang: str) -> str:
    from app.config import settings

    try:
        if settings.GOOGLE_CLOUD_API_KEY:
            cloud_response = requests.get(
                "https://translation.googleapis.com/language/translate/v2",
                params={
                    "key": settings.GOOGLE_CLOUD_API_KEY,
                    "q": text[:5000],
                    "source": source_lang,
                    "target": target_lang,
                    "format": "text",
                },
                timeout=60,
            )
            if cloud_response.status_code == 200:
                data = cloud_response.json()
                if data.get("data", {}).get("translations"):
                    logger.info(
                        f"Translated using Google Translate {source_lang} -> {target_lang}"
                    )
                    return data["data"]["translations"][0]["translatedText"]
    except Exception as e:
        logger.warning(f"Google Translate failed: {e}")

    logger.error("Translation failed - no API available")
    return text


def _batch_translate(texts: list[str], source_lang: str, target_lang: str) -> list[str]:
    """Translate a list of texts in one API call using concurrent requests."""
    from app.config import settings
    from concurrent.futures import ThreadPoolExecutor, as_completed

    if not texts:
        return texts

    # Try Google Translate batch (up to 128 strings per request)
    try:
        if settings.GOOGLE_CLOUD_API_KEY:
            import urllib.parse
            # Build POST body with repeated q= params
            body = "&".join(
                f"q={urllib.parse.quote(str(t)[:5000])}" for t in texts
            ) + f"&source={source_lang}&target={target_lang}&format=text"

            response = requests.post(
                f"https://translation.googleapis.com/language/translate/v2?key={settings.GOOGLE_CLOUD_API_KEY}",
                data=body,
                headers={"Content-Type": "application/x-www-form-urlencoded"},
                timeout=60,
            )
            if response.status_code == 200:
                translations = response.json().get("data", {}).get("translations", [])
                if len(translations) == len(texts):
                    logger.info(f"Batch translated {len(texts)} texts {source_lang}->{target_lang}")
                    return [t["translatedText"] for t in translations]
    except Exception as e:
        logger.warning(f"Batch translate failed: {e}")

    # Fallback: concurrent individual requests
    results = [None] * len(texts)
    with ThreadPoolExecutor(max_workers=10) as executor:
        futures = {executor.submit(translate_chunk, t, source_lang, target_lang): i for i, t in enumerate(texts)}
        for future in as_completed(futures):
            results[futures[future]] = future.result()
    return results


def _translate_excel_json(original_text: str, source_lang: str, target_lang: str) -> str:
    """Translate cell values and sheet names in batches."""
    try:
        data = json.loads(original_text)
    except Exception:
        return original_text

    sheets = data.get("sheets") or {}
    sheet_names = data.get("sheet_names") or list(sheets.keys())

    # Collect all unique non-empty strings across all sheets + sheet names
    all_texts = list(sheet_names)
    for sheet_name in sheet_names:
        for row in sheets.get(sheet_name, []):
            for cell in row:
                if cell and str(cell).strip():
                    all_texts.append(str(cell))

    # Batch translate in chunks of 100
    BATCH = 100
    translated_map = {}
    for i in range(0, len(all_texts), BATCH):
        batch = all_texts[i:i + BATCH]
        results = _batch_translate(batch, source_lang, target_lang)
        for orig, trans in zip(batch, results):
            translated_map[orig] = trans

    translated_sheet_names = [translated_map.get(n, n) for n in sheet_names]

    translated_sheets = {}
    for sheet_name, trans_name in zip(sheet_names, translated_sheet_names):
        rows = sheets.get(sheet_name, [])
        translated_rows = [
            [translated_map.get(str(cell), cell) if cell and str(cell).strip() else cell for cell in row]
            for row in rows
        ]
        translated_sheets[trans_name] = translated_rows

    return json.dumps({
        "sheet_names": sheet_names,
        "translated_sheet_names": translated_sheet_names,
        "sheets": translated_sheets,
    })


class TranslationTask(Task):
    autoretry_for = (Exception,)
    retry_backoff = 5
    retry_backoff_max = 60
    max_retries = 3


@celery_app.task(bind=True, base=TranslationTask)
def translate_content(
    self,
    translation_id: str,
    original_text: str,
    language_id: int,
    source_language_id: int | None = None,
):
    logger.info(f"Translation task started: {translation_id}")

    db = SessionLocal()
    try:
        from app.models import Translation, TranslationJob, Language

        target_language = db.query(Language).filter(Language.id == language_id).first()
        if not target_language:
            raise ValueError(f"Invalid language_id: {language_id}")

        target_lang_code = target_language.libretranslate_code or target_language.code

        source_language = None
        if source_language_id:
            source_language = (
                db.query(Language).filter(Language.id == source_language_id).first()
            )

        source_lang_code = (
            source_language.libretranslate_code if source_language else "en"
        )

        logger.info(f"Translating from {source_lang_code} to {target_lang_code}")

        # Check if this is an exam (JSON with sheets structure) — translate cells, not raw JSON
        is_excel = False
        try:
            parsed = json.loads(original_text)
            if isinstance(parsed, dict) and ("sheets" in parsed or "makaratasi" in parsed):
                is_excel = True
        except Exception:
            pass

        if is_excel:
            translated_text = _translate_excel_json(original_text, source_lang_code, target_lang_code)
            chunks = [original_text]  # for chunk_count
        else:
            # Translate line by line to preserve structure
            lines = original_text.split("\n")
            non_empty_lines = [(i, l) for i, l in enumerate(lines) if l.strip()]
            chunks = non_empty_lines  # for chunk_count

            # Batch translate only non-empty lines
            BATCH = 100
            translated_map = {}
            texts = [l for _, l in non_empty_lines]
            for i in range(0, len(texts), BATCH):
                batch = texts[i:i+BATCH]
                results = _batch_translate(batch, source_lang_code, target_lang_code)
                for (idx, _), trans in zip(non_empty_lines[i:i+BATCH], results):
                    translated_map[idx] = trans

            # Rebuild with same line structure
            translated_lines = [translated_map.get(i, l) for i, l in enumerate(lines)]
            translated_text = "\n".join(translated_lines)

        translation = (
            db.query(Translation)
            .filter(Translation.id == uuid.UUID(translation_id))
            .first()
        )
        if translation:
            translation.translated_text = translated_text
            translation.status = "done"
            translation.translation_engine = "google_translate"
            translation.chunk_count = len(chunks)
            db.commit()

        job = (
            db.query(TranslationJob)
            .filter(TranslationJob.celery_task_id == self.request.id)
            .first()
        )
        if job:
            job.completed_at = datetime.utcnow()
            db.commit()

        # Pre-generate cached PDF for .docx books so downloads are instant
        if translation and translation.content_type == "book" and not is_excel:
            try:
                from app.models import Book
                book = db.query(Book).filter(Book.id == translation.content_id).first()
                if book and book.file_path and book.file_path.endswith(".docx"):
                    cached_pdf_key = book.file_path.replace(".docx", f"_translated_{translation.language_id}.pdf")
                    cached_pdf_path = f"{settings.STORAGE_ROOT}/{cached_pdf_key}"
                    if not os.path.exists(cached_pdf_path):
                        from reportlab.lib.pagesizes import A4
                        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
                        from reportlab.lib.styles import ParagraphStyle
                        from reportlab.lib.units import inch
                        from reportlab.pdfbase import pdfmetrics
                        from reportlab.pdfbase.ttfonts import TTFont
                        import io as _io, fitz as _fitz

                        from docx import Document

                        pdfmetrics.registerFont(TTFont("DejaVu", "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"))
                        pdfmetrics.registerFont(TTFont("DejaVu-Bold", "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf"))

                        title_style = ParagraphStyle("t",  fontName="DejaVu-Bold", fontSize=16, spaceAfter=10, leading=20, alignment=1)
                        h1_style    = ParagraphStyle("h1", fontName="DejaVu-Bold", fontSize=14, spaceAfter=8,  leading=18, spaceBefore=10)
                        h2_style    = ParagraphStyle("h2", fontName="DejaVu-Bold", fontSize=12, spaceAfter=6,  leading=16, spaceBefore=6)
                        h3_style    = ParagraphStyle("h3", fontName="DejaVu-Bold", fontSize=11, spaceAfter=4,  leading=14, spaceBefore=4)
                        list_style  = ParagraphStyle("li", fontName="DejaVu",      fontSize=10, spaceAfter=3,  leading=13, leftIndent=20)
                        body_style  = ParagraphStyle("b",  fontName="DejaVu",      fontSize=10, spaceAfter=4,  leading=14)

                        # Build translation map from translated_text lines
                        trans_lines = [l for l in translated_text.split("\n") if l.strip()]
                        trans_iter = iter(trans_lines)

                        # Read original docx paragraphs for structure
                        with open(f"{settings.STORAGE_ROOT}/{book.file_path}", "rb") as f:
                            orig_doc = Document(f)

                        non_empty = [p for p in orig_doc.paragraphs if p.text.strip()]
                        # Last 5 non-empty paras = last page (keep in English)
                        last_page_texts = {p.text.strip() for p in non_empty[-5:]}
                        # First 4 pages worth of paragraphs (keep in English)
                        first_content_page = book.first_content_page or 5
                        skip_count = first_content_page - 1  # pages to skip
                        # Estimate paragraphs per page (~15) to find skip boundary
                        skip_para_count = skip_count * 15
                        skip_texts = {p.text.strip() for p in non_empty[:skip_para_count] if p.text.strip()}

                        buf = _io.BytesIO()
                        pdf = SimpleDocTemplate(buf, pagesize=A4,
                            leftMargin=0.75*inch, rightMargin=0.75*inch,
                            topMargin=0.75*inch, bottomMargin=0.75*inch)
                        story = []

                        for para in orig_doc.paragraphs:
                            orig_text = para.text.strip()

                            # Handle inline images — embed as image in PDF
                            has_image = (
                                para._element.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}inline') or
                                para._element.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}anchor')
                            )
                            if has_image:
                                try:
                                    from reportlab.platypus import Image as RLImage
                                    DML = "http://schemas.openxmlformats.org/drawingml/2006/main"
                                    REL = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
                                    for blip in para._element.findall(f'.//{{{DML}}}blip'):
                                        rId = blip.get(f'{{{REL}}}embed')
                                        if rId and rId in orig_doc.part.rels:
                                            img_part = orig_doc.part.rels[rId].target_part
                                            img_buf = _io.BytesIO(img_part.blob)
                                            img = RLImage(img_buf, width=4*inch, height=3*inch)
                                            story.append(img)
                                            story.append(Spacer(1, 0.1*inch))
                                except Exception:
                                    pass
                                if not orig_text:
                                    continue

                            if not orig_text:
                                story.append(Spacer(1, 0.08*inch))
                                continue

                            sname = para.style.name.lower() if para.style else ""

                            # Keep first N pages and last page in original English
                            if orig_text in last_page_texts or orig_text in skip_texts:
                                text = orig_text
                            else:
                                text = next(trans_iter, orig_text)

                            safe = text.replace("&","&amp;").replace("<","&lt;").replace(">","&gt;")

                            if "title" in sname:
                                story.append(Paragraph(safe, title_style))
                            elif "heading 1" in sname:
                                story.append(Paragraph(safe, h1_style))
                            elif "heading 2" in sname:
                                story.append(Paragraph(safe, h2_style))
                            elif "heading 3" in sname:
                                story.append(Paragraph(safe, h3_style))
                            elif "list" in sname:
                                story.append(Paragraph(f"• {safe}", list_style))
                            else:
                                story.append(Paragraph(safe, body_style))

                        pdf.build(story)

                        # Prepend first N pages as images from original PDF
                        source_pdf_path = f"{settings.STORAGE_ROOT}/{book.file_path.replace('.docx', '.pdf')}"
                        body_doc = _fitz.open("pdf", buf.getvalue())

                        if os.path.exists(source_pdf_path):
                            src_pdf = _fitz.open(source_pdf_path)
                            pages_to_prepend = min(skip_count, len(src_pdf))
                            for page_num in range(pages_to_prepend - 1, -1, -1):
                                pix = src_pdf[page_num].get_pixmap(matrix=_fitz.Matrix(1.5, 1.5))
                                new_page = body_doc.new_page(0, width=pix.width/1.5, height=pix.height/1.5)
                                new_page.insert_image(new_page.rect, pixmap=pix)
                        elif os.path.exists(f"{settings.STORAGE_ROOT}/{book.file_path.replace('.docx', '_cover.png')}"):
                            # Fallback: just cover image
                            cover_pix = _fitz.Pixmap(f"{settings.STORAGE_ROOT}/{book.file_path.replace('.docx', '_cover.png')}")
                            cover_page = body_doc.new_page(0, width=cover_pix.width/2, height=cover_pix.height/2)
                            cover_page.insert_image(cover_page.rect, pixmap=cover_pix)

                        final_buf = _io.BytesIO()
                        body_doc.save(final_buf)
                        final_bytes = final_buf.getvalue()

                        with open(cached_pdf_path, "wb") as f:
                            f.write(final_bytes)
                        logger.info(f"Cached PDF generated: {cached_pdf_path}")
            except Exception as e:
                logger.warning(f"Pre-generation of cached PDF failed: {e}")

        logger.info(f"Translation task completed: {translation_id}")
        return {"status": "success", "translation_id": translation_id}

    except Exception as e:
        logger.error(f"Translation task failed: {e}")

        try:
            translation = (
                db.query(Translation)
                .filter(Translation.id == uuid.UUID(translation_id))
                .first()
            )
            if translation:
                translation.status = "failed"
                db.commit()

            job = (
                db.query(TranslationJob)
                .filter(TranslationJob.celery_task_id == self.request.id)
                .first()
            )
            if job:
                job.error_message = str(e)
                db.commit()
        except Exception:
            pass

        raise
    finally:
        db.close()
